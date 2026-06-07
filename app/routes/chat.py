"""Chat routes — send messages, stream responses, upload files."""

from __future__ import annotations

import json
import uuid
import logging
import re
from pathlib import Path

from fastapi import APIRouter, HTTPException, UploadFile, File, Form, Query, Depends
from fastapi.responses import StreamingResponse, Response

from app.core.config import settings
from app.core.auth import get_current_user
from app.schemas.chat import ChatRequest, ChatResponse, FileUploadResponse
from app.db.store import store
from app.services import llm
from app.services.analysis import run_pipeline
from app.services.document_parser import get_parser
from app.services.export_service import generate_docx, generate_pdf

logger = logging.getLogger(__name__)

router = APIRouter(prefix="/api/chat", tags=["chat"])

# ─── File Export ─────────────────────────────────────────────

@router.get("/export/{conversation_id}")
async def export_conversation(
    conversation_id: str,
    format: str = Query("docx", enum=["pdf", "docx"]),
    user_id: str = Depends(get_current_user)
):
    """Export the conversation report as PDF or DOCX."""
    conv = await store.get_conversation(conversation_id, user_id=user_id)
    if not conv:
        raise HTTPException(status_code=404, detail="Conversation not found")
        
    if format == "docx":
        file_stream = await generate_docx(conversation_id)
        if not file_stream:
            raise HTTPException(status_code=404, detail="Could not generate DOCX. Ensure a report exists.")
            
        filename = f"Fortress_AI_Report_{conversation_id[:8]}.docx"
        return Response(
            content=file_stream.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    
    elif format == "pdf":
        file_stream = await generate_pdf(conversation_id)
        if not file_stream:
            raise HTTPException(status_code=404, detail="Could not generate PDF. Ensure a report exists.")
            
        filename = f"Fortress_AI_Report_{conversation_id[:8]}.pdf"
        return Response(
            content=file_stream.getvalue(),
            media_type="application/pdf",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )

    raise HTTPException(status_code=400, detail="Unsupported format")

# ─── System prompt for conversational chat ───────────────────

FORTRESS_SYSTEM_PROMPT = """You are Fortress AI, a legal contract risk analysis assistant.

Rules:
- Return only the final answer for the user. Never output chain-of-thought, hidden reasoning, or internal deliberation.
- Never output tags or phrases like <think>, </think>, "Thinking Process", "Here's a thinking process", "Step 1", or similar.
- Keep responses concise, professional, and actionable.
- Use clean Markdown when helpful. Do not wrap tables in code fences.
- Do not provide definitive legal advice; include a short disclaimer when legal conclusions are requested.
- If the user asks for a downloadable report, provide:
    [Download DOCX Report](/api/chat/export/{conversation_id}?format=docx)
"""


def _sanitize_assistant_output(text: str) -> str:
    """Remove chain-of-thought artifacts if the model leaks them."""
    if not text:
        return text

    cleaned = text

    if "</think>" in cleaned:
        tail = cleaned.rsplit("</think>", 1)[-1].strip()
        if tail:
            cleaned = tail

    cleaned = re.sub(r"<think>.*?</think>", "", cleaned, flags=re.IGNORECASE | re.DOTALL)

    # Fallback: if a reasoning header remains, keep only the final paragraph.
    if re.search(r"thinking\s*process|here'?s\s+a\s+thinking\s+process", cleaned, flags=re.IGNORECASE):
        blocks = [b.strip() for b in re.split(r"\n\s*\n", cleaned) if b.strip()]
        if blocks:
            cleaned = blocks[-1]

    return cleaned.strip()


# ─── Chat (non-streaming) ───────────────────────────────────

@router.post("", response_model=ChatResponse)
async def send_message(req: ChatRequest, user_id: str = Depends(get_current_user)):
    """Send a message and get a complete response."""

    # Auto-create conversation if none provided
    conv_id = req.conversation_id
    if not conv_id:
        conv_id = uuid.uuid4().hex[:12]
        title = req.message[:50].strip() or "New Analysis"
        await store.create_conversation(
            id=conv_id,
            user_id=user_id,
            title=title,
            contract_type=req.contract_type.value if req.contract_type else None,
            user_type=req.user_type.value if req.user_type else None,
        )
    else:
        conv = await store.get_conversation(conv_id, user_id=user_id)
        if not conv:
            raise HTTPException(status_code=404, detail="Conversation not found")

    # Save user message
    user_msg_id = uuid.uuid4().hex[:12]
    await store.add_message(
        conversation_id=conv_id,
        user_id=user_id,
        message_id=user_msg_id,
        role="user",
        content=req.message,
    )

    # Build conversation history for multi-turn
    history = await store.get_messages(conv_id, user_id=user_id)
    messages = [{"role": m.role.value, "content": m.content} for m in history]

    # Generate response
    system = FORTRESS_SYSTEM_PROMPT.replace("{conversation_id}", conv_id)
    if req.user_type:
        system += f"\n\nThe user's role is: {req.user_type.value}."
    if req.contract_type:
        system += f"\nThey are analyzing a: {req.contract_type.value} contract."

    response_text = await llm.generate_with_history(
        messages=messages,
        system_prompt=system,
    )
    response_text = _sanitize_assistant_output(response_text)

    # Save assistant message
    assistant_msg_id = uuid.uuid4().hex[:12]
    assistant_msg = await store.add_message(
        conversation_id=conv_id,
        user_id=user_id,
        message_id=assistant_msg_id,
        role="assistant",
        content=response_text,
    )

    if not assistant_msg:
        raise HTTPException(status_code=500, detail="Failed to save assistant message")

    return ChatResponse(message=assistant_msg, conversation_id=conv_id)


# ─── Chat (SSE streaming) ───────────────────────────────────

@router.post("/stream")
async def stream_message(req: ChatRequest, user_id: str = Depends(get_current_user)):
    """Send a message and get a streamed SSE response."""

    # Auto-create conversation if none provided
    conv_id = req.conversation_id
    if not conv_id:
        conv_id = uuid.uuid4().hex[:12]
        title = req.message[:50].strip() or "New Analysis"
        await store.create_conversation(
            id=conv_id,
            user_id=user_id,
            title=title,
            contract_type=req.contract_type.value if req.contract_type else None,
            user_type=req.user_type.value if req.user_type else None,
        )
    else:
        conv = await store.get_conversation(conv_id, user_id=user_id)
        if not conv:
            raise HTTPException(status_code=404, detail="Conversation not found")

    # Save user message
    user_msg_id = uuid.uuid4().hex[:12]
    await store.add_message(
        conversation_id=conv_id,
        user_id=user_id,
        message_id=user_msg_id,
        role="user",
        content=req.message,
    )

    # Build history
    history = await store.get_messages(conv_id, user_id=user_id)
    messages = [{"role": m.role.value, "content": m.content} for m in history]

    system = FORTRESS_SYSTEM_PROMPT.replace("{conversation_id}", conv_id)
    if req.user_type:
        system += f"\n\nThe user's role is: {req.user_type.value}."
    if req.contract_type:
        system += f"\nThey are analyzing a: {req.contract_type.value} contract."

    assistant_msg_id = uuid.uuid4().hex[:12]

    async def event_stream():
        full_response = []

        # Send conversation_id and message_id first
        yield f"data: {json.dumps({'event': 'start', 'conversation_id': conv_id, 'message_id': assistant_msg_id})}\n\n"

        try:
            async for chunk in llm.stream_with_history(
                messages=messages,
                system_prompt=system,
            ):
                full_response.append(chunk)
                yield f"data: {json.dumps({'event': 'chunk', 'content': chunk})}\n\n"
        except Exception as e:
            logger.error(f"Stream error: {e}")
            yield f"data: {json.dumps({'event': 'error', 'message': str(e)})}\n\n"
            return

        # Save the full response
        complete_text = _sanitize_assistant_output("".join(full_response))
        await store.add_message(
            conversation_id=conv_id,
            user_id=user_id,
            message_id=assistant_msg_id,
            role="assistant",
            content=complete_text,
        )

        yield f"data: {json.dumps({'event': 'done', 'conversation_id': conv_id})}\n\n"

    return StreamingResponse(event_stream(), media_type="text/event-stream")


# ─── Audit Pipeline (SSE streaming) ─────────────────────────

@router.post("/audit")
async def stream_audit(req: ChatRequest, user_id: str = Depends(get_current_user)):
    """Run the full multi-step audit pipeline with SSE progress events."""

    conv_id = req.conversation_id
    if not conv_id:
        conv_id = uuid.uuid4().hex[:12]
        title = req.message[:50].strip() or "Contract Analysis"
        await store.create_conversation(
            id=conv_id,
            user_id=user_id,
            title=title,
            contract_type=req.contract_type.value if req.contract_type else None,
            user_type=req.user_type.value if req.user_type else None,
        )
    else:
        conv = await store.get_conversation(conv_id, user_id=user_id)
        if not conv:
            raise HTTPException(status_code=404, detail="Conversation not found")

    # Save user message
    user_msg_id = uuid.uuid4().hex[:12]
    await store.add_message(
        conversation_id=conv_id,
        user_id=user_id,
        message_id=user_msg_id,
        role="user",
        content=req.message,
    )

    history = await store.get_messages(conv_id, user_id=user_id)

    # Prefer uploaded document text from system messages when available.
    uploaded_segments: list[str] = []
    for msg in history:
        if msg.role.value != "system":
            continue
        if not msg.content.startswith("[File uploaded:"):
            continue

        marker = "\n\n"
        marker_pos = msg.content.find(marker)
        if marker_pos == -1:
            continue

        uploaded_segments.append(msg.content[marker_pos + len(marker):])

    # Use only the latest uploaded document segment to avoid unbounded growth
    # when users upload multiple revisions in one conversation.
    uploaded_text = ""
    for seg in reversed(uploaded_segments):
        if seg.strip():
            uploaded_text = seg
            break
    if uploaded_text.endswith("..."):
        uploaded_text = uploaded_text[:-3].rstrip()

    pipeline_text = req.message
    if uploaded_text.strip():
        pipeline_text = uploaded_text

    # Keep LLM calls bounded to reduce timeout risk on very large documents.
    MAX_PIPELINE_CHARS = 30000
    if len(pipeline_text) > MAX_PIPELINE_CHARS:
        pipeline_text = pipeline_text[:MAX_PIPELINE_CHARS]

    # Guardrail: prevent running the full pipeline on tiny non-document prompts.
    if not uploaded_text.strip() and len((req.message or "").strip()) < 300:
        async def too_short_stream():
            yield f"data: {json.dumps({'event': 'start', 'conversation_id': conv_id})}\n\n"
            yield f"data: {json.dumps({'event': 'error', 'data': {'message': 'No document text found. Please upload a document first or paste the full contract text before running audit.'}})}\n\n"

        return StreamingResponse(too_short_stream(), media_type="text/event-stream")

    # Zero-token precheck
    if uploaded_text.strip():
        is_contract_like, precheck_reason = _precheck_contract_document(uploaded_text)
        if not is_contract_like:
            async def non_contract_stream():
                yield f"data: {json.dumps({'event': 'start', 'conversation_id': conv_id})}\n\n"
                yield f"data: {json.dumps({'event': 'error', 'data': {'message': precheck_reason}})}\n\n"

            return StreamingResponse(non_contract_stream(), media_type="text/event-stream")

    logger.info(
        "Audit input | conversation_id=%s | uploaded_segments=%d | uploaded_chars=%d | message_chars=%d",
        conv_id,
        len(uploaded_segments),
        len(uploaded_text),
        len(req.message or ""),
    )

    async def pipeline_stream():
        yield f"data: {json.dumps({'event': 'start', 'conversation_id': conv_id})}\n\n"

        async for event in run_pipeline(
            text=pipeline_text,
            contract_type=req.contract_type.value if req.contract_type else None,
            user_type=req.user_type.value if req.user_type else None,
        ):
            yield event

    return StreamingResponse(pipeline_stream(), media_type="text/event-stream")


# ─── File Upload ─────────────────────────────────────────────

@router.options("/upload")
async def upload_options():
    """Handle CORS preflight for upload endpoint."""
    return Response(status_code=200)

@router.post("/upload", response_model=FileUploadResponse)
async def upload_file(
    file: UploadFile = File(...),
    conversation_id: str = Form(...),
    user_id: str = Depends(get_current_user)
):
    """Upload a document (PDF, DOCX, TXT) to a conversation."""

    allowed_types = {
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "text/plain",
    }
    if file.content_type not in allowed_types:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported file type: {file.content_type}. Allowed: PDF, DOCX, TXT",
        )

    contents = await file.read()
    max_bytes = settings.MAX_UPLOAD_SIZE_MB * 1024 * 1024
    if len(contents) > max_bytes:
        raise HTTPException(
            status_code=400,
            detail=f"File too large. Maximum: {settings.MAX_UPLOAD_SIZE_MB}MB",
        )

    conv = await store.get_conversation(conversation_id, user_id=user_id)
    if conv is None:
        raise HTTPException(status_code=404, detail="Conversation not found")

    file_id = uuid.uuid4().hex[:12]
    ext = Path(file.filename).suffix if file.filename else ".bin"
    save_path = settings.upload_path / f"{file_id}{ext}"
    save_path.write_bytes(contents)

    extracted_text = await _extract_text(contents, file.content_type, file.filename)
    
    # NEW: MCP + RAG Integration for PDFs
    mcp_metadata = {}
    if file.content_type == "application/pdf":
        try:
            from app.services.mcp_integration import get_integration_service
            integration = get_integration_service()
            
            result = await integration.process_uploaded_document(
                file_id=file_id,
                pdf_bytes=contents,
                filename=file.filename or "document.pdf",
                conversation_id=conversation_id
            )
            
            if result["success"]:
                mcp_metadata = {
                    "mcp_registered": True,
                    "chunks_created": result["chunks_created"],
                    "indexed": result["indexed"],
                    "sections": result["sections"],
                    "pages": result["pages"]
                }
                logger.info(f"MCP integration successful for {file_id}: {mcp_metadata}")
            else:
                logger.warning(f"MCP integration failed for {file_id}: {result.get('error')}")
        except Exception as e:
            logger.warning(f"MCP integration error (non-fatal): {e}")
            # Continue without MCP - existing flow still works

    attachment = {
        "id": file_id,
        "name": file.filename or "document",
        "size": len(contents),
        "type": file.content_type or "application/octet-stream",
        **mcp_metadata  # Include MCP metadata if available
    }

    await store.add_message(
        conversation_id=conversation_id,
        user_id=user_id,
        message_id=uuid.uuid4().hex[:12],
        role="system",
        content=f"[File uploaded: {file.filename}]\n\n{extracted_text}",
        attachment=attachment,
    )

    return FileUploadResponse(
        id=file_id,
        name=file.filename or "document",
        size=len(contents),
        type=file.content_type or "application/octet-stream",
        conversation_id=conversation_id,
    )


async def _extract_text(contents: bytes, content_type: str, filename: str | None) -> str:
    """Extract text from uploaded file based on type."""
    try:
        if content_type == "text/plain":
            return contents.decode("utf-8", errors="replace")

        elif content_type == "application/pdf":
            try:
                parser = get_parser()
                parsed = parser.parse_pdf(contents)

                output_parts: list[str] = []
                title = parsed.metadata.get("title")
                if title:
                    output_parts.append(f"# {title}\n")
                output_parts.append(parsed.text)

                logger.info(
                    "PDF parsed: %s | pages=%d | blocks=%d | tables=%d | sections=%d",
                    filename,
                    parsed.page_count,
                    len(parsed.blocks),
                    len(parsed.tables),
                    len(parsed.sections),
                )

                return "\n".join(part for part in output_parts if part)
            except ImportError:
                logger.warning("PDF parser dependencies not available")
                return f"[PDF file: {filename}]"

        elif "wordprocessingml" in (content_type or ""):
            try:
                import docx
                import io
                doc = docx.Document(io.BytesIO(contents))
                return "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            except ImportError:
                return f"[DOCX file: {filename} — install python-docx for text extraction]"

        return f"[Unsupported file format: {content_type}]"

    except Exception as e:
        logger.error(f"Text extraction failed: {e}")
        return f"[Error extracting text from {filename}: {str(e)}]"


def _precheck_contract_document(text: str) -> tuple[bool, str]:
    """Cheap heuristic gate to block obvious non-contract files."""
    sample = (text or "").strip().lower()[:12000]
    if len(sample) < 120:
        return (
            False,
            "Uploaded text is too short. Please upload the full contract document.",
        )

    contract_markers = [
        "agreement", "contract", "nda", "msa", "sow", "lease", "vendor",
        "terms and conditions", "obligations", "termination", "confidential",
        "governing law", "liability", "indemn", "warranty", "payment"
    ]
    
    contract_hits = sum(1 for marker in contract_markers if marker in sample)
    numbered_sections = len(re.findall(r"(?m)^\s*(\d+(?:\.\d+)*)\s+[A-Z][A-Za-z ]{2,}$", text[:12000]))

    if contract_hits < 2 and numbered_sections == 0:
        return (
            False,
            "Could not detect contract structure. Please upload a contract/agreement document.",
        )

    return True, "ok"
